"""주택임대차 계약서 생성 테스트 스크립트"""
import sys
sys.stdout.reconfigure(encoding='utf-8')

import asyncio
import json
from pathlib import Path
from docx import Document

# 스키마 로드
schema_path = Path(__file__).parent / "data" / "storage" / "documents" / "lease_contract_input_schema.json"
with open(schema_path, 'r', encoding='utf-8') as f:
    schema = json.load(f)

# 예시 데이터 가져오기
test_data_1 = schema["샘플_데이터_예시"]["예시1_전세"]
test_data_2 = schema["샘플_데이터_예시"]["예시2_월세"]

print("=" * 80)
print("주택임대차 계약서 생성 테스트")
print("=" * 80)

async def test_contract_generation():
    """계약서 생성 테스트"""
    from app.service_agent.tools.lease_contract_generator_tool import LeaseContractGeneratorTool

    tool = LeaseContractGeneratorTool()

    print("\n[테스트 1] 전세 계약서 생성")
    print("-" * 80)
    result1 = await tool.execute(**test_data_1)

    if result1["status"] == "success":
        print(f"✅ 생성 성공: {result1['docx_path']}")
        print(f"\n{result1['content']}\n")

        # 생성된 파일 검증
        verify_generated_file(result1['docx_path'], test_data_1)
    else:
        print(f"❌ 생성 실패: {result1.get('error') or result1.get('message')}")

    print("\n" + "=" * 80)
    print("[테스트 2] 월세 계약서 생성")
    print("-" * 80)
    result2 = await tool.execute(**test_data_2)

    if result2["status"] == "success":
        print(f"✅ 생성 성공: {result2['docx_path']}")
        print(f"\n{result2['content']}\n")

        # 생성된 파일 검증
        verify_generated_file(result2['docx_path'], test_data_2)
    else:
        print(f"❌ 생성 실패: {result2.get('error') or result2.get('message')}")


def verify_generated_file(docx_path: str, expected_data: dict):
    """생성된 DOCX 파일 검증"""
    print("\n📋 생성된 파일 검증")
    print("-" * 80)

    doc = Document(docx_path)

    # 전체 텍스트 추출
    all_text = "\n".join([p.text for p in doc.paragraphs])

    # 테이블 텍스트 추출
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += "\n" + cell.text

    # 필수 필드 검증
    checks = []

    if "address_road" in expected_data:
        found = expected_data["address_road"] in all_text
        checks.append(("주소", expected_data["address_road"], found))

    if "deposit" in expected_data:
        # 콤마 제거하고 검색
        deposit_clean = expected_data["deposit"].replace(",", "")
        found = deposit_clean in all_text or expected_data["deposit"] in all_text
        checks.append(("보증금", expected_data["deposit"], found))

    if "lessor_name" in expected_data:
        found = expected_data["lessor_name"] in all_text
        checks.append(("임대인 이름", expected_data["lessor_name"], found))

    if "lessee_name" in expected_data:
        found = expected_data["lessee_name"] in all_text
        checks.append(("임차인 이름", expected_data["lessee_name"], found))

    if "monthly_rent" in expected_data:
        monthly_rent_clean = expected_data["monthly_rent"].replace(",", "")
        found = monthly_rent_clean in all_text or expected_data["monthly_rent"] in all_text
        checks.append(("월세", expected_data["monthly_rent"], found))

    # 결과 출력
    all_passed = True
    for field_name, expected_value, found in checks:
        status = "✅" if found else "❌"
        print(f"{status} {field_name}: {expected_value} {'발견됨' if found else '발견 안됨'}")
        if not found:
            all_passed = False

    if all_passed:
        print("\n🎉 모든 필드가 정상적으로 채워졌습니다!")
    else:
        print("\n⚠️  일부 필드가 누락되었습니다. 코드 수정이 필요합니다.")

    print("-" * 80)


if __name__ == "__main__":
    asyncio.run(test_contract_generation())
