"""
.doc  .docx  
     
"""
import os
import sys
import shutil
from pathlib import Path
from typing import List, Optional, Tuple
import logging
import traceback

#  
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

class DocConverter:
    """DOC to DOCX """

    def __init__(self, output_dir: str = None):
        self.output_dir = Path(output_dir) if output_dir else Path.cwd() / "converted"
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.conversion_stats = {
            'success': [],
            'failed': [],
            'skipped': []
        }

    def convert_with_win32com(self, doc_path: Path) -> Optional[Path]:
        """Win32COM   (Windows )"""
        try:
            import win32com.client

            # Word Application 
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            try:
                #   
                doc_path_str = str(doc_path.absolute())
                output_path = self.output_dir / f"{doc_path.stem}.docx"
                output_path_str = str(output_path.absolute())

                #  
                doc = word.Documents.Open(doc_path_str, ReadOnly=True)

                # DOCX  (16 = wdFormatXMLDocument)
                doc.SaveAs2(output_path_str, FileFormat=16)
                doc.Close()

                logger.info(f"Converted with Win32COM: {doc_path.name}")
                return output_path

            finally:
                word.Quit()

        except Exception as e:
            logger.warning(f"Win32COM  : {e}")
            return None

    def convert_with_doc2docx(self, doc_path: Path) -> Optional[Path]:
        """doc2docx   """
        try:
            from doc2docx import convert

            output_path = self.output_dir / f"{doc_path.stem}.docx"
            convert(str(doc_path), str(output_path))

            logger.info(f"Converted with doc2docx: {doc_path.name}")
            return output_path

        except ImportError:
            logger.warning("doc2docx not installed. Install with: pip install doc2docx")
            return None
        except Exception as e:
            logger.warning(f"doc2docx  : {e}")
            return None

    def convert_with_python_docx(self, doc_path: Path) -> Optional[Path]:
        """python-docx    ()"""
        try:
            import docx

            #  .doc    
            doc = docx.Document(str(doc_path))
            output_path = self.output_dir / f"{doc_path.stem}.docx"
            doc.save(str(output_path))

            logger.info(f"Converted with python-docx: {doc_path.name}")
            return output_path

        except Exception as e:
            logger.warning(f"python-docx  : {e}")
            return None

    def extract_text_fallback(self, doc_path: Path) -> Optional[Path]:
        """   DOCX  ( )"""
        try:
            import docx
            import subprocess

            # antiword catdoc   
            text = None

            # antiword 
            try:
                result = subprocess.run(
                    ['antiword', str(doc_path)],
                    capture_output=True,
                    text=True,
                    encoding='utf-8'
                )
                if result.returncode == 0:
                    text = result.stdout
            except:
                pass

            if not text:
                #    ( )
                with open(doc_path, 'rb') as f:
                    content = f.read()
                    # UTF-16 LE  
                    try:
                        text = content.decode('utf-16-le', errors='ignore')
                        #    
                        text = ''.join(c for c in text if c.isprintable() or c in '\n\r\t')
                    except:
                        text = content.decode('utf-8', errors='ignore')

            if text and len(text) > 100:  #    
                #  DOCX 
                doc = docx.Document()

                #  
                doc.add_heading(f'Converted from: {doc_path.name}', 0)
                doc.add_paragraph('Warning: Original format may be lost.')
                doc.add_paragraph('')

                #    
                for line in text.split('\n'):
                    if line.strip():
                        doc.add_paragraph(line)

                output_path = self.output_dir / f"{doc_path.stem}_text_only.docx"
                doc.save(str(output_path))

                logger.info(f"Text extracted to DOCX: {doc_path.name}")
                return output_path

        except Exception as e:
            logger.error(f"Text extraction failed: {e}")
            return None

    def convert_single(self, doc_path: Path) -> Tuple[bool, Optional[Path]]:
        """   (   )"""

        #  .docx  
        if doc_path.suffix.lower() == '.docx':
            output_path = self.output_dir / doc_path.name
            shutil.copy2(doc_path, output_path)
            logger.info(f"Copied existing DOCX: {doc_path.name}")
            return True, output_path

        #    
        conversion_methods = [
            ('Win32COM', self.convert_with_win32com),
            ('doc2docx', self.convert_with_doc2docx),
            ('python-docx', self.convert_with_python_docx),
            ('Text Extraction', self.extract_text_fallback)
        ]

        for method_name, method in conversion_methods:
            logger.info(f"Trying {method_name}...")
            result = method(doc_path)
            if result and result.exists():
                self.conversion_stats['success'].append({
                    'file': doc_path.name,
                    'method': method_name,
                    'output': str(result)
                })
                return True, result

        #   
        self.conversion_stats['failed'].append({
            'file': doc_path.name,
            'error': 'All conversion methods failed'
        })
        logger.error(f"Failed to convert: {doc_path.name}")
        return False, None

    def convert_batch(self, doc_paths: List[Path]) -> List[Path]:
        """ """
        converted_files = []

        print(f"\n{'='*60}")
        print(f"Converting {len(doc_paths)} DOC files to DOCX")
        print(f"{'='*60}\n")

        for i, doc_path in enumerate(doc_paths, 1):
            print(f"\n[{i}/{len(doc_paths)}] Processing: {doc_path.name}")
            print("-" * 40)

            success, output_path = self.convert_single(doc_path)

            if success and output_path:
                converted_files.append(output_path)
                print(f"Success: {output_path.name}")
            else:
                print(f"Failed: {doc_path.name}")

        self.print_summary()
        return converted_files

    def print_summary(self):
        """  """
        print(f"\n{'='*60}")
        print("Conversion Summary")
        print(f"{'='*60}")

        print(f"Success: {len(self.conversion_stats['success'])} files")
        for item in self.conversion_stats['success']:
            print(f"   - {item['file']} ({item['method']})")

        if self.conversion_stats['failed']:
            print(f"\nFailed: {len(self.conversion_stats['failed'])} files")
            for item in self.conversion_stats['failed']:
                print(f"   - {item['file']}: {item['error']}")

        if self.conversion_stats['skipped']:
            print(f"\nSkipped: {len(self.conversion_stats['skipped'])} files")
            for item in self.conversion_stats['skipped']:
                print(f"   - {item}")

        print(f"\nOutput directory: {self.output_dir}")


def main():
    """  """
    #    
    source_dir = Path(r"C:\kdy\Projects\holmesnyangz\beta_v001\backend\data\___\4. \1.  (·)")
    output_dir = Path(r"C:\kdy\Projects\holmesnyangz\beta_v001\backend\data\___\vector_maker\data\raw")

    #  
    converter = DocConverter(output_dir=output_dir)

    # .doc  
    doc_files = list(source_dir.glob("*.doc"))

    if not doc_files:
        print("No .doc files found!")
        return

    print(f"Found {len(doc_files)} .doc files:")
    for f in doc_files:
        print(f"  - {f.name}")

    #   
    converted_files = converter.convert_batch(doc_files)

    print(f"\nConversion complete!")
    print(f"Converted files saved to: {output_dir}")

    return converted_files


if __name__ == "__main__":
    # doc2docx  
    try:
        import doc2docx
    except ImportError:
        print("Installing doc2docx...")
        os.system("pip install doc2docx")

    converted_files = main()