"""
문서 전처리 모듈
- 표 데이터 정규화
- 텍스트 클리닝
- 구조 보존 전처리
"""
import re
from typing import List, Dict, Any, Tuple
import docx
from pathlib import Path
import pandas as pd
import logging

logger = logging.getLogger(__name__)

class DocumentPreprocessor:
    """문서 전처리기"""
    
    def __init__(self):
        self.cleaning_patterns = {
            'whitespace': r'\s+',
            'special_chars': r'[^\w\s가-힣.,;:()《》\[\]「」『』【】\-/]',
            'repeated_dots': r'\.{3,}',
            'parenthesis_numbers': r'\(\d+\)',
            'bullet_points': r'^[\-\*\•]\s*',
        }
    
    def preprocess(self, file_path: str, structure_info: Dict = None) -> Dict[str, Any]:
        """문서 전처리 메인 함수"""
        doc = docx.Document(file_path)
        
        processed_data = {
            'tables': [],
            'paragraphs': [],
            'structured_content': [],
            'metadata': {}
        }
        
        # 표 처리
        for table in doc.tables:
            processed_table = self._process_table(table)
            if processed_table:
                processed_data['tables'].append(processed_table)
                processed_data['structured_content'].append({
                    'type': 'table',
                    'content': processed_table
                })
        
        # 텍스트 처리
        for para in doc.paragraphs:
            processed_text = self._process_paragraph(para)
            if processed_text:
                processed_data['paragraphs'].append(processed_text)
                processed_data['structured_content'].append({
                    'type': 'text',
                    'content': processed_text
                })
        
        return processed_data
    
    def _process_table(self, table) -> Dict[str, Any]:
        """표 데이터 처리"""
        processed_table = {
            'headers': [],
            'rows': [],
            'formatted_text': '',
            'structured_data': []
        }
        
        # 헤더 추출
        if len(table.rows) > 0:
            headers = []
            for cell in table.rows[0].cells:
                header = self._clean_text(cell.text)
                headers.append(header)
            processed_table['headers'] = headers
        
        # 데이터 행 처리
        for row_idx, row in enumerate(table.rows[1:]):
            row_data = []
            row_dict = {}
            
            for col_idx, cell in enumerate(row.cells):
                cell_text = self._clean_text(cell.text)
                row_data.append(cell_text)
                
                if col_idx < len(processed_table['headers']):
                    header = processed_table['headers'][col_idx]
                    row_dict[header] = cell_text
            
            processed_table['rows'].append(row_data)
            processed_table['structured_data'].append(row_dict)
        
        # 표를 텍스트로 변환 (RAG용)
        processed_table['formatted_text'] = self._format_table_as_text(
            processed_table['headers'],
            processed_table['rows']
        )
        
        return processed_table
    
    def _process_paragraph(self, paragraph) -> Dict[str, Any]:
        """단락 처리"""
        text = paragraph.text.strip()
        if not text:
            return None
        
        # 텍스트 클리닝
        cleaned_text = self._clean_text(text)
        
        # 구조 정보 추출
        structure_info = {
            'original': text,
            'cleaned': cleaned_text,
            'style': paragraph.style.name if paragraph.style else None,
            'is_heading': 'Heading' in (paragraph.style.name or ''),
            'has_numbering': self._has_numbering(text),
            'law_references': self._extract_law_references(text)
        }
        
        return structure_info
    
    def _clean_text(self, text: str) -> str:
        """텍스트 클리닝"""
        if not text:
            return ''
        
        # 기본 클리닝
        text = text.strip()
        
        # 과도한 공백 제거
        text = re.sub(self.cleaning_patterns['whitespace'], ' ', text)
        
        # 반복된 점 정규화
        text = re.sub(self.cleaning_patterns['repeated_dots'], '...', text)
        
        # 불필요한 특수문자 제거 (법률 문서 특수문자는 보존)
        # text = re.sub(self.cleaning_patterns['special_chars'], '', text)
        
        return text
    
    def _format_table_as_text(self, headers: List[str], rows: List[List[str]]) -> str:
        """표를 구조화된 텍스트로 변환"""
        formatted_lines = []
        
        # 헤더가 있는 경우
        if headers:
            for row in rows:
                row_text = []
                for i, (header, cell) in enumerate(zip(headers, row)):
                    if header and cell:
                        # 법률 문서 표 특별 포맷팅
                        if '구분' in header:
                            row_text.append(f"[{header}: {cell}]")
                        elif '법률' in header or '법령' in header:
                            row_text.append(f"{header}: {cell}")
                        elif '내용' in header:
                            row_text.append(f"- {header}: {cell}")
                        else:
                            row_text.append(f"{header}: {cell}")
                
                formatted_lines.append(' '.join(row_text))
        else:
            # 헤더가 없는 경우
            for row in rows:
                formatted_lines.append(' | '.join(row))
        
        return '\n'.join(formatted_lines)
    
    def _has_numbering(self, text: str) -> bool:
        """번호 매기기 여부 확인"""
        numbering_patterns = [
            r'^\d+\.',
            r'^\d+\)',
            r'^[가-하]\.',
            r'^[ㄱ-ㅎ]\.',
            r'^제\d+조',
            r'^제\d+항'
        ]
        
        for pattern in numbering_patterns:
            if re.match(pattern, text.strip()):
                return True
        return False
    
    def _extract_law_references(self, text: str) -> List[str]:
        """법률 참조 추출"""
        references = []
        
        # 법률 번호
        law_numbers = re.findall(r'법률\s*제?\d+호', text)
        references.extend(law_numbers)
        
        # 조항 참조
        articles = re.findall(r'제\d+조(?:\s*제\d+항)?', text)
        references.extend(articles)
        
        # 시행령/시행규칙
        regulations = re.findall(r'시행령\s*제?\d+호|시행규칙\s*제?\d+호', text)
        references.extend(regulations)
        
        return references
    
    def save_processed_data(self, processed_data: Dict, output_path: str):
        """전처리된 데이터 저장"""
        output_path = Path(output_path)
        
        # 텍스트 파일로 저장
        with open(output_path.with_suffix('.txt'), 'w', encoding='utf-8') as f:
            for item in processed_data['structured_content']:
                if item['type'] == 'table':
                    f.write("\n[TABLE]\n")
                    f.write(item['content']['formatted_text'])
                    f.write("\n[/TABLE]\n")
                elif item['type'] == 'text':
                    f.write(item['content']['cleaned'])
                    f.write("\n")
        
        # 구조화된 데이터를 JSON으로도 저장
        import json
        json_path = output_path.with_suffix('.json')
        
        # DataFrame을 dict로 변환하여 JSON 직렬화 가능하게 만들기
        json_data = {
            'tables': [
                {
                    'headers': t['headers'],
                    'rows': t['rows'],
                    'formatted_text': t['formatted_text']
                } for t in processed_data['tables']
            ],
            'paragraphs': [
                {
                    'cleaned': p['cleaned'],
                    'style': p['style'],
                    'is_heading': p['is_heading']
                } for p in processed_data['paragraphs']
            ]
        }
        
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(json_data, f, ensure_ascii=False, indent=2)
        
        logger.info(f"Processed data saved to {output_path}")
    
    def print_preprocessing_summary(self, processed_data: Dict):
        """전처리 결과 요약 출력"""
        print("\n" + "="*50)
        print("🔧 전처리 결과")
        print("="*50)
        print(f"표 개수: {len(processed_data['tables'])}")
        print(f"단락 개수: {len(processed_data['paragraphs'])}")
        
        # 표 데이터 샘플
        if processed_data['tables']:
            print("\n📊 표 데이터 샘플:")
            table = processed_data['tables'][0]
            print(f"  헤더: {table['headers']}")
            if table['rows']:
                print(f"  첫 번째 행: {table['rows'][0]}")
        
        # 텍스트 데이터 샘플
        if processed_data['paragraphs']:
            print("\n📝 텍스트 데이터 샘플:")
            para = processed_data['paragraphs'][0]
            print(f"  {para['cleaned'][:100]}...")
        
        print("="*50)
