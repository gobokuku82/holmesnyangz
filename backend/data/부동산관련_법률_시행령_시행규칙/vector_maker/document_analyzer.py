"""
문서 구조 파악 모듈
- 표 구조 감지
- 계층 구조 분석
- 메타데이터 추출
"""
import re
from typing import Dict, List, Any, Optional
from dataclasses import dataclass
from pathlib import Path
import docx
from docx.table import Table
import logging

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

@dataclass
class DocumentStructure:
    """문서 구조 정보"""
    doc_type: str  # 문서 타입 (table, mixed, text)
    has_tables: bool
    table_count: int
    sections: List[Dict[str, Any]]
    metadata: Dict[str, Any]
    hierarchy: Dict[str, Any]

class DocumentAnalyzer:
    """문서 구조 분석기"""
    
    def __init__(self):
        self.structure_patterns = {
            '구분': r'구분|분류|범주|카테고리',
            '법률': r'법률|법령|조항|규정|조문',
            '내용': r'내용|설명|규정사항|주요\s*내용',
            '조항': r'제\d+조|제\d+항|제\d+호',
            '날짜': r'\d{4}년\s*\d{1,2}월\s*\d{1,2}일|\d{8}',
        }
    
    def analyze(self, file_path: str) -> DocumentStructure:
        """문서 전체 구조 분석"""
        file_path = Path(file_path)
        
        if file_path.suffix.lower() in ['.docx', '.doc']:
            return self._analyze_word(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_path.suffix}")
    
    def _analyze_word(self, file_path: Path) -> DocumentStructure:
        """Word 문서 구조 분석"""
        doc = docx.Document(str(file_path))
        
        # 표 분석
        tables_info = self._analyze_tables(doc.tables)
        
        # 텍스트 구조 분석
        text_structure = self._analyze_text_structure(doc.paragraphs)
        
        # 메타데이터 추출
        metadata = self._extract_metadata(doc)
        
        # 문서 타입 결정
        doc_type = self._determine_doc_type(tables_info, text_structure)
        
        return DocumentStructure(
            doc_type=doc_type,
            has_tables=len(tables_info) > 0,
            table_count=len(tables_info),
            sections=self._merge_sections(tables_info, text_structure),
            metadata=metadata,
            hierarchy=self._build_hierarchy(tables_info, text_structure)
        )
    
    def _analyze_tables(self, tables: List[Table]) -> List[Dict[str, Any]]:
        """표 구조 분석"""
        tables_info = []
        
        for idx, table in enumerate(tables):
            table_info = {
                'type': 'table',
                'index': idx,
                'rows': len(table.rows),
                'cols': len(table.columns),
                'headers': [],
                'content_preview': [],
                'structure_type': None  # legal_table, comparison_table, etc.
            }
            
            # 헤더 추출
            if len(table.rows) > 0:
                headers = []
                for cell in table.rows[0].cells:
                    header_text = cell.text.strip()
                    headers.append(header_text)
                    
                    # 법률 문서 표 타입 판별
                    if any(re.search(pattern, header_text) 
                          for pattern in self.structure_patterns.values()):
                        table_info['structure_type'] = 'legal_table'
                
                table_info['headers'] = headers
            
            # 내용 미리보기 (처음 3행)
            for row_idx, row in enumerate(table.rows[1:4]):
                row_content = [cell.text.strip()[:50] for cell in row.cells]
                table_info['content_preview'].append(row_content)
            
            tables_info.append(table_info)
        
        return tables_info
    
    def _analyze_text_structure(self, paragraphs) -> List[Dict[str, Any]]:
        """텍스트 구조 분석"""
        text_structure = []
        current_section = None
        
        for para in paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 제목/섹션 감지
            if para.style.name and 'Heading' in para.style.name:
                if current_section:
                    text_structure.append(current_section)
                current_section = {
                    'type': 'section',
                    'title': text,
                    'level': int(para.style.name[-1]) if para.style.name[-1].isdigit() else 0,
                    'content': []
                }
            elif current_section:
                current_section['content'].append(text)
            else:
                text_structure.append({
                    'type': 'paragraph',
                    'content': text
                })
        
        if current_section:
            text_structure.append(current_section)
        
        return text_structure
    
    def _extract_metadata(self, doc) -> Dict[str, Any]:
        """메타데이터 추출"""
        metadata = {}
        
        # 문서 속성
        if doc.core_properties:
            metadata['title'] = doc.core_properties.title or ''
            metadata['author'] = doc.core_properties.author or ''
            metadata['created'] = str(doc.core_properties.created) if doc.core_properties.created else ''
            metadata['modified'] = str(doc.core_properties.modified) if doc.core_properties.modified else ''
        
        # 법률 관련 메타데이터 추출 (텍스트에서)
        all_text = '\n'.join([p.text for p in doc.paragraphs[:10]])  # 처음 10개 단락
        
        # 법령 번호 추출
        law_numbers = re.findall(r'제\d+호|법률\s*제\d+호', all_text)
        if law_numbers:
            metadata['law_numbers'] = law_numbers
        
        # 날짜 추출
        dates = re.findall(r'\d{4}년\s*\d{1,2}월\s*\d{1,2}일', all_text)
        if dates:
            metadata['dates'] = dates
        
        return metadata
    
    def _determine_doc_type(self, tables_info: List, text_structure: List) -> str:
        """문서 타입 결정"""
        if not tables_info:
            return 'text'
        elif len(tables_info) > len(text_structure):
            return 'table_heavy'
        else:
            return 'mixed'
    
    def _merge_sections(self, tables_info: List, text_structure: List) -> List[Dict]:
        """표와 텍스트 섹션 병합"""
        # 간단한 병합 - 실제로는 위치 정보를 기반으로 정교하게 병합
        sections = []
        sections.extend(tables_info)
        sections.extend(text_structure)
        return sections
    
    def _build_hierarchy(self, tables_info: List, text_structure: List) -> Dict:
        """계층 구조 구축"""
        hierarchy = {
            'root': {
                'tables': [t['index'] for t in tables_info],
                'sections': [s.get('title', 'untitled') for s in text_structure if s['type'] == 'section']
            }
        }
        return hierarchy
    
    def print_structure_summary(self, structure: DocumentStructure):
        """구조 요약 출력"""
        print("\n" + "="*50)
        print("📄 문서 구조 분석 결과")
        print("="*50)
        print(f"문서 타입: {structure.doc_type}")
        print(f"표 개수: {structure.table_count}")
        
        if structure.has_tables:
            print("\n📊 표 정보:")
            for section in structure.sections:
                if section.get('type') == 'table':
                    print(f"  - 표 {section['index']}: {section['rows']}행 x {section['cols']}열")
                    print(f"    헤더: {', '.join(section['headers'][:3])}")
        
        if structure.metadata:
            print("\n📌 메타데이터:")
            for key, value in structure.metadata.items():
                if value:
                    print(f"  - {key}: {value}")
        
        print("="*50)
